import argparse
from pathlib import Path

import mido
from docx import Document
from docx.shared import Pt

from music21 import converter
from docx2pdf import convert as docx_to_pdf

# ================= CONFIG =================

NOTE_WIDTH = 4
MAX_CHARS_PER_LINE = 64
MAX_NOTES_PER_LINE = MAX_CHARS_PER_LINE // NOTE_WIDTH

FONT = "Courier New"
FONT_SIZE = 11

NOTE_NAMES = ["C","C#","D","D#","E","F","F#","G","G#","A","A#","B"]

# Alto sax fingering (simplified)
FINGERINGS = {
    "C":  [1,1,1,1,1,1,1],
    "D":  [1,1,1,1,1,1,0],
    "E":  [1,1,1,1,1,0,0],
    "F":  [1,1,1,1,0,0,0],
    "G":  [1,1,1,0,0,0,0],
    "A":  [1,1,0,0,0,0,0],
    "B":  [1,0,0,0,0,0,0],
}

# ================= PARSERS =================

def midi_notes(path):
    midi = mido.MidiFile(path)
    notes = []

    for track in midi.tracks:          # <-- ALL tracks
        for msg in track:
            if msg.type == "note_on" and msg.velocity > 0:
                name = NOTE_NAMES[msg.note % 12].replace("#", "")
                notes.append(name)

    return notes


def mcsz_notes(path):
    score = converter.parse(path)
    notes = []
    for n in score.recurse().notes:
        if n.isNote:
            notes.append(n.pitch.name.replace("#",""))
    return notes


def docx_notes(path):
    doc = Document(path)
    notes = []
    for p in doc.paragraphs:
        for tok in p.text.split():
            if tok.upper() in NOTE_NAMES:
                notes.append(tok.upper())
    return notes

# ================= RENDERING =================


def render_fingering_rows(notes):
    rows = [""] * 7  # 7 vertical buttons

    for n in notes:
        fingering = FINGERINGS.get(n, [0]*7)
        for i in range(7):
            symbol = "●" if fingering[i] else "○"
            rows[i] += symbol.ljust(NOTE_WIDTH)

    return rows

# ================= OUTPUT =================

def write_docx(notes, out_path, fingering=False):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(FONT_SIZE)

    doc.add_heading("Alto Sax Guide", level=1)

    for i in range(0, len(notes), MAX_NOTES_PER_LINE):
        chunk = notes[i:i+MAX_NOTES_PER_LINE]

    for i in range(0, len(notes), MAX_NOTES_PER_LINE):
        chunk = notes[i:i+MAX_NOTES_PER_LINE]
    
        if fingering:
            fingering_rows = render_fingering_rows(chunk)
            for row in fingering_rows:
                doc.add_paragraph(row)
    
        note_line = "".join(n.ljust(NOTE_WIDTH) for n in chunk)
        lyric_line = "".join("".ljust(NOTE_WIDTH) for _ in chunk)
    
        doc.add_paragraph(note_line)
        doc.add_paragraph(lyric_line)
        doc.add_paragraph("")

    doc.save(out_path)

# ================= MAIN =================

def main():
    p = argparse.ArgumentParser()
    p.add_argument("input")
    p.add_argument("--out", required=True)
    p.add_argument("--fingering", action="store_true")
    p.add_argument("--pdf", action="store_true")

    args = p.parse_args()
    inp = Path(args.input)

    if inp.suffix == ".mid":
        notes = midi_notes(inp)
    elif inp.suffix == ".mcsz":
        notes = mcsz_notes(inp)
    elif inp.suffix == ".docx":
        notes = docx_notes(inp)
    else:
        raise ValueError("Unsupported input")

    write_docx(notes, Path(args.out), args.fingering)

    if args.pdf:
        docx_to_pdf(args.out)

if __name__ == "__main__":
    main()